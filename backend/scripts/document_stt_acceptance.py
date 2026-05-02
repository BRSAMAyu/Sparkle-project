import asyncio
import json
import mimetypes
import os
import subprocess
import tempfile
import time
import wave
from pathlib import Path

import requests
import websockets
from docx import Document

from _acceptance_common import login_with_requests


BASE_URL = "http://127.0.0.1:8080/api/v1"
AUTH_BASE_URL = "http://127.0.0.1:8000/api/v1"
WS_URL = "ws://127.0.0.1:8080/ws/stt"
USERNAME = os.getenv("LOCAL_SMOKE_USERNAME", "chat_test")
PASSWORD = os.getenv("LOCAL_SMOKE_PASSWORD", "Chat123456")
REQUEST_TIMEOUT_SECONDS = 180
VOICE_TEXT = (
    "Sparkle speech acceptance test. "
    "This recording verifies the speech to text pipeline across the gateway and backend."
)


def _request(method: str, path: str, *, token: str | None = None, expected_status: int = 200, **kwargs):
    headers = kwargs.pop("headers", {})
    if token:
        headers = {"Authorization": f"Bearer {token}", **headers}

    response = requests.request(
        method,
        f"{BASE_URL}{path}",
        headers=headers,
        timeout=REQUEST_TIMEOUT_SECONDS,
        **kwargs,
    )
    if response.status_code != expected_status:
        raise RuntimeError(
            f"{method} {path} expected {expected_status}, got {response.status_code}: {response.text[:600]}"
        )
    return response

def _build_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("长期计划与知识沉淀示例文档", level=1)
    doc.add_paragraph("这是一份用于验收文档清洗链路的样本文档。")
    doc.add_paragraph("它包含可直接抽取的文字层，适合验证 DOCX 解析、清洗、摘要与结果轮询。")
    doc.add_paragraph("核心目标：确保移动端文档清洗工具在真实网关路径下可用。")
    doc.save(path)


def _run_document_cleaning(token: str, workdir: Path) -> dict:
    started_at = time.monotonic()
    docx_path = workdir / "document_cleaning_acceptance.docx"
    _build_docx(docx_path)

    mime_type = mimetypes.guess_type(docx_path.name)[0] or "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    with docx_path.open("rb") as handle:
        response = _request(
            "POST",
            "/documents/clean",
            token=token,
            files={"file": (docx_path.name, handle, mime_type)},
            data={"options": json.dumps({"enable_ocr": False, "ocr_engine": "zhipu"})},
        )

    payload = response.json()
    task_id = payload["task_id"]

    deadline = time.monotonic() + 90
    last_status = None
    while time.monotonic() < deadline:
        status = _request("GET", f"/documents/clean/{task_id}", token=token).json()
        last_status = status
        if status.get("status") in {"completed", "failed", "error"}:
            break
        time.sleep(1)

    if not last_status:
        raise RuntimeError("Document cleaning did not return status")
    if last_status.get("status") != "completed":
        raise RuntimeError(f"Document cleaning failed: {json.dumps(last_status, ensure_ascii=False)}")

    result = last_status.get("result") or {}
    summary = (result.get("summary") or "").strip()
    if not summary:
        raise RuntimeError(f"Document cleaning returned empty summary: {json.dumps(last_status, ensure_ascii=False)}")

    return {
        "task_id": task_id,
        "mode": result.get("mode"),
        "summary_preview": summary[:120],
        "char_count": result.get("char_count"),
        "elapsed_seconds": round(time.monotonic() - started_at, 2),
    }


def _generate_wav_with_tts(workdir: Path) -> Path:
    aiff_path = workdir / "stt_acceptance.aiff"
    wav_path = workdir / "stt_acceptance.wav"

    subprocess.run(
        ["say", "-o", str(aiff_path), VOICE_TEXT],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    subprocess.run(
        ["ffmpeg", "-y", "-i", str(aiff_path), "-ar", "16000", "-ac", "1", str(wav_path)],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return wav_path


def _run_http_transcribe(token: str, wav_path: Path) -> dict:
    started_at = time.monotonic()
    with wav_path.open("rb") as handle:
        response = _request(
            "POST",
            "/stt/transcribe",
            token=token,
            files={"file": (wav_path.name, handle, "audio/wav")},
            data={"language": "en-US"},
        )

    payload = response.json()
    if payload.get("error"):
        raise RuntimeError(f"HTTP STT failed: {json.dumps(payload, ensure_ascii=False)}")
    text = (payload.get("text") or "").strip()
    enhanced = (payload.get("enhanced_text") or "").strip()
    if not text:
        raise RuntimeError(f"HTTP STT returned empty text: {json.dumps(payload, ensure_ascii=False)}")

    return {
        "text": text,
        "enhanced_text": enhanced,
        "elapsed_seconds": round(time.monotonic() - started_at, 2),
    }


async def _run_ws_transcribe(token: str, wav_path: Path) -> dict:
    started_at = time.monotonic()
    with wave.open(str(wav_path), "rb") as wav_file:
        audio_frames = wav_file.readframes(wav_file.getnframes())

    uri = f"{WS_URL}?token={token}"
    transcripts: list[str] = []
    completed = False

    async with websockets.connect(uri, max_size=None) as websocket:
        chunk_size = 3200
        for start in range(0, len(audio_frames), chunk_size):
            await websocket.send(audio_frames[start:start + chunk_size])
            await asyncio.sleep(0.02)

        await websocket.send("STOP")

        deadline = time.monotonic() + 20
        while time.monotonic() < deadline:
            try:
                message = await asyncio.wait_for(websocket.recv(), timeout=5)
            except TimeoutError as exc:
                raise RuntimeError("Timed out waiting for WebSocket transcription") from exc

            payload = json.loads(message)
            if payload.get("type") == "transcription":
                text = (payload.get("text") or "").strip()
                if text:
                    transcripts.append(text)
            elif payload.get("type") == "status" and payload.get("content") == "completed":
                completed = True
                break
            elif payload.get("type") == "error":
                raise RuntimeError(f"WebSocket STT failed: {json.dumps(payload, ensure_ascii=False)}")

    if not completed:
        raise RuntimeError("WebSocket STT did not complete cleanly")
    if not transcripts:
        raise RuntimeError("WebSocket STT returned no transcript")

    return {
        "text": transcripts[-1],
        "segments": len(transcripts),
        "elapsed_seconds": round(time.monotonic() - started_at, 2),
    }


def main() -> None:
    token = login_with_requests(
        auth_base_url=AUTH_BASE_URL,
        username=USERNAME,
        password=PASSWORD,
        timeout_seconds=REQUEST_TIMEOUT_SECONDS,
    )

    with tempfile.TemporaryDirectory(prefix="document-stt-acceptance-") as tmpdir:
        workdir = Path(tmpdir)

        document_result = _run_document_cleaning(token, workdir)
        wav_path = _generate_wav_with_tts(workdir)
        ws_stt_result = asyncio.run(_run_ws_transcribe(token, wav_path))
        http_stt_result = _run_http_transcribe(token, wav_path)

    result = {
        "document": document_result,
        "stt_http": {
            "text_preview": http_stt_result["text"][:120],
            "enhanced_preview": http_stt_result["enhanced_text"][:120],
        },
        "stt_ws": ws_stt_result,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    print("ALL_OK")


if __name__ == "__main__":
    main()
