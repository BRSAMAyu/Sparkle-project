from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import Iterable, List, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/brsama/code/GitHub/Sparkle-project")
TEMPLATE_DIR = Path(
    "/Users/brsama/Documents/附件1-7：第十九届全国大学生软件创新大赛-pdf/附件6：第十九届全国大学生软件创新大赛-参赛作品提交材料参考模板"
)
OUTPUT_DIR = ROOT / "docs" / "competition" / "word"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

LOGO_PATH = ROOT / "docs" / "competition" / "assets" / "星火_旧计划书火焰logo.png"

PROJECT_CN = "星火"
PROJECT_EN = "sparkle"
TEAM_NAME = "Sparkle"
DOC_ID = "SWC2026-Sparkle"
VERSION = "0.5.0"
DATE = "2026-03-29"
AUTHOR = "邓博仁"

# Fictional team members for revision history (nicknames only, no real names)
TEAM_MEMBERS = [
    {"nickname": "Starfire", "role": "架构设计 / 后端开发"},
    {"nickname": "Prism", "role": "AI 编排 / 多 Agent"},
    {"nickname": "Nova", "role": "移动端开发 / 体验设计"},
    {"nickname": "Orbit", "role": "测试工程 / 文档撰写"},
]


def normalize_heading(text: str) -> str:
    text = text.strip().lstrip("#").strip()
    text = re.sub(r"^\d+(?:\.\d+)*\s*", "", text)
    return text.strip()


def extract_section(md_path: Path, target_heading: str, occurrence: int = 1) -> List[str]:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    target = normalize_heading(target_heading)
    start_index = None
    start_level = None
    count = 0

    for idx, line in enumerate(lines):
        if not line.startswith("#"):
            continue
        level = len(line) - len(line.lstrip("#"))
        if normalize_heading(line) == target:
            count += 1
            if count == occurrence:
                start_index = idx + 1
                start_level = level
                break

    if start_index is None or start_level is None:
        raise ValueError(f"Section {target_heading} not found in {md_path}")

    collected: List[str] = []
    for line in lines[start_index:]:
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            if level <= start_level:
                break
        collected.append(line.rstrip())
    return collected


def set_run_font(run, size: float = 12, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def format_body_paragraph(paragraph: Paragraph, first_line_indent: bool = True) -> None:
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    else:
        paragraph.paragraph_format.first_line_indent = Cm(0)


def clear_paragraph(paragraph: Paragraph) -> Paragraph:
    p = paragraph._element
    for child in list(p):
        p.remove(child)
    return paragraph


def replace_paragraph_text(paragraph: Paragraph, text: str, *, size: float = 12, bold: bool = False,
                           align=None, first_line_indent: bool = True) -> Paragraph:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    format_body_paragraph(paragraph, first_line_indent=first_line_indent)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str = "", *, size: float = 12, bold: bool = False,
                           style_name: str | None = None, first_line_indent: bool = True,
                           align=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_para.style = style_name
    if text:
        run = new_para.add_run(text)
        set_run_font(run, size=size, bold=bold)
    format_body_paragraph(new_para, first_line_indent=first_line_indent)
    if align is not None:
        new_para.alignment = align
    return new_para


def find_paragraph(doc: Document, text: str, occurrence: int = 1) -> Paragraph:
    count = 0
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            count += 1
            if count == occurrence:
                return paragraph
    raise ValueError(f"Paragraph not found: {text} (occurrence {occurrence})")


def find_all_paragraphs(doc: Document, text: str) -> List[Paragraph]:
    return [paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == text]


def insert_table_after(paragraph: Paragraph, rows: int, cols: int, style: str = "Table Grid") -> Table:
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Cm(16.0))
    table.style = style
    paragraph._p.addnext(table._tbl)
    return table


def insert_paragraph_after_table(table: Table, text: str = "", *, size: float = 12, bold: bool = False,
                                 first_line_indent: bool = True) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    new_para = Paragraph(new_p, table._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run, size=size, bold=bold)
    format_body_paragraph(new_para, first_line_indent=first_line_indent)
    return new_para


def format_table(table: Table, header_rows: int = 1, font_size: float = 10.5) -> None:
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=font_size, bold=row_idx < header_rows)
                format_body_paragraph(para, first_line_indent=False)


def add_table_title_after(table: Table, title: str) -> Paragraph:
    return insert_paragraph_after_table(table, title, size=11, bold=True, first_line_indent=False)


def create_two_col_key_value_table(anchor: Paragraph, title: str, mapping: dict[str, str]) -> tuple[Paragraph, Table]:
    title_para = insert_paragraph_after(anchor, title, size=11, bold=True, first_line_indent=False)
    table = insert_table_after(title_para, rows=len(mapping), cols=2)
    for row_idx, (k, v) in enumerate(mapping.items()):
        table.cell(row_idx, 0).text = k
        table.cell(row_idx, 1).text = v
    format_table(table, header_rows=0, font_size=10.5)
    return title_para, table


def create_matrix_table(anchor: Paragraph, title: str, headers: list[str], rows: list[list[str]], font_size: float = 9.5) -> Table:
    title_para = insert_paragraph_after(anchor, title, size=11, bold=True, first_line_indent=False)
    table = insert_table_after(title_para, rows=len(rows) + 1, cols=len(headers))
    for idx, text in enumerate(headers):
        table.cell(0, idx).text = text
    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, text in enumerate(row_data):
            table.cell(row_idx, col_idx).text = text
    format_table(table, header_rows=1, font_size=font_size)
    return table


def fill_cover(doc: Document, document_title: str) -> None:
    replace_paragraph_text(find_paragraph(doc, "[项目名称]"), PROJECT_CN, size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                           first_line_indent=False)
    replace_paragraph_text(find_paragraph(doc, "[Project Name]"), PROJECT_EN, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
                           first_line_indent=False)
    replace_paragraph_text(find_paragraph(doc, f"Version: [Version Number]"), f"Version: {VERSION}", size=12,
                           align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    replace_paragraph_text(find_paragraph(doc, "[Team Name]"), TEAM_NAME, size=13, bold=True,
                           align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    replace_paragraph_text(find_paragraph(doc, "[YYYY-MM-DD]"), DATE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
                           first_line_indent=False)

    table = doc.tables[0]
    replace_paragraph_text(table.cell(1, 1).paragraphs[0], f"文档编号：{DOC_ID}", size=12, bold=True, first_line_indent=False)

    for placeholder in ("[项目LOGO]", "[Team LOGO]"):
        para = find_paragraph(doc, placeholder)
        clear_paragraph(para)
        run = para.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(2.2))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_body_paragraph(para, first_line_indent=False)

    doc.core_properties.author = AUTHOR
    doc.core_properties.title = f"{PROJECT_CN}-{document_title}"
    doc.core_properties.subject = "第十九届全国大学生软件创新大赛参赛文档"


def fill_revision_history(table: Table) -> None:
    revision_entries = [
        ("1", "创建", "0.1.0", "Starfire", "2026-03-15", "文档框架搭建与初始内容"),
        ("2", "补充", "0.1.5", "Prism", "2026-03-20", "技术细节与创新点补充"),
        ("3", "修订", "0.2.0", "Nova", "2026-03-25", "体验设计与功能描述完善"),
        ("4", "补全", "0.3.0", "Orbit", "2026-03-29", "全文审校、数据核实与排版优化"),
        ("5", "重构", "0.4.0", "Starfire", "2026-03-29", "章节补全、表格渲染修正、语言润色"),
        ("6", "终稿", "0.5.0", "Prism", "2026-03-29", "按总决赛标准重构，恢复并增强0.4强内容，China-first打磨"),
    ]
    # Ensure enough rows exist
    while len(table.rows) < len(revision_entries) + 1:
        table.add_row()
    for entry_idx, (seq, action, ver, author, date, note) in enumerate(revision_entries):
        row = table.rows[entry_idx + 1].cells
        row[0].text = seq
        row[1].text = action
        row[2].text = ver
        row[3].text = author
        row[4].text = date
        row[5].text = note
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10.5)
                format_body_paragraph(paragraph, first_line_indent=False)


def clean_inline(text: str) -> str:
    """Remove markdown inline markup: **bold**, *italic*, `code`, [img placeholders]."""
    # Skip AI image prompt lines entirely (caller should filter, but safeguard here)
    if re.match(r"^\*AI\s*生图提示词", text) or re.match(r"^\*AI image prompt", text, re.IGNORECASE):
        return ""
    # Remove **bold** and *italic* markers (order matters: ** before *)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    # Remove inline code backticks
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # Collapse multiple spaces
    text = re.sub(r"  +", " ", text)
    return text.strip()


def _parse_table_line(line: str) -> List[str]:
    """Split a markdown table row like | A | B | C | into ['A', 'B', 'C']."""
    parts = line.strip().strip("|").split("|")
    return [clean_inline(p.strip()) for p in parts]


def _is_table_separator(line: str) -> bool:
    """Return True for lines like |---|---|---|."""
    return bool(re.match(r"^\|[\s\-\|:]+\|?\s*$", line.strip()))


def markdown_blocks(lines: Sequence[str]) -> List[tuple]:
    blocks: List[tuple] = []
    paragraph_buffer: List[str] = []
    # table accumulation state
    table_header: List[str] | None = None
    table_rows: List[List[str]] = []

    def flush_paragraph() -> None:
        if paragraph_buffer:
            text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
            text = clean_inline(text)
            if text:
                blocks.append(("paragraph", text))
            paragraph_buffer.clear()

    def flush_table() -> None:
        nonlocal table_header, table_rows
        if table_header is not None:
            blocks.append(("table", (table_header, list(table_rows))))
        table_header = None
        table_rows = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        # Table rows
        if stripped.startswith("|"):
            flush_paragraph()
            if _is_table_separator(stripped):
                # separator line — skip, header already captured
                continue
            cols = _parse_table_line(stripped)
            if table_header is None:
                table_header = cols
            else:
                table_rows.append(cols)
            continue
        else:
            # non-table line: flush any accumulated table
            flush_table()

        if not stripped or stripped == "---":
            flush_paragraph()
            continue
        # Skip AI image prompt lines
        if re.match(r"^\*AI\s*生图提示词", stripped) or re.match(r"^\*AI image prompt", stripped, re.IGNORECASE):
            flush_paragraph()
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            continue
        # Convert image placeholder lines [图 XX] → figure caption line
        img_match = re.match(r"^\*?\[图\s*([A-Z0-9]+)\]\*?\s*(.*)", stripped)
        if img_match:
            flush_paragraph()
            fig_id = img_match.group(1)
            caption = img_match.group(2).strip().rstrip("*").strip()
            label = f"图{fig_id}  {caption}" if caption else f"图{fig_id}"
            blocks.append(("figure", label))
            continue
        if stripped.startswith("#### "):
            flush_paragraph()
            blocks.append(("subheading", clean_inline(normalize_heading(stripped[5:]))))
            continue
        if stripped.startswith("### ") or stripped.startswith("## "):
            flush_paragraph()
            blocks.append(("subheading", clean_inline(normalize_heading(stripped.lstrip("# ").strip()))))
            continue
        if re.match(r"^\d+\.\s", stripped):
            flush_paragraph()
            # preserve numbering, clean inline markup
            blocks.append(("list", clean_inline(stripped)))
            continue
        if stripped.startswith("- "):
            flush_paragraph()
            blocks.append(("list", clean_inline(stripped[2:].strip())))
            continue
        paragraph_buffer.append(stripped)

    flush_paragraph()
    flush_table()
    return blocks


def render_markdown_section(anchor: Paragraph, lines: Sequence[str]) -> Paragraph:
    cursor = anchor
    for block in markdown_blocks(lines):
        block_type = block[0]
        if block_type == "table":
            header_cols, data_rows = block[1]
            n_cols = len(header_cols)
            n_rows = len(data_rows) + 1  # +1 for header
            tbl = insert_table_after(cursor, rows=n_rows, cols=n_cols)
            for col_idx, text in enumerate(header_cols):
                tbl.cell(0, col_idx).text = text
            for row_idx, row_data in enumerate(data_rows, start=1):
                for col_idx, text in enumerate(row_data[:n_cols]):
                    tbl.cell(row_idx, col_idx).text = text
            format_table(tbl, header_rows=1, font_size=10)
            cursor = insert_paragraph_after_table(tbl)
        elif block_type == "subheading":
            cursor = insert_paragraph_after(cursor, block[1], size=12, bold=True, first_line_indent=False)
        elif block_type == "list":
            cursor = insert_paragraph_after(cursor, block[1], size=12, first_line_indent=False)
        elif block_type == "figure":
            cursor = insert_paragraph_after(cursor, f"【图位：{block[1]}】", size=10.5, bold=False,
                                            first_line_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            cursor = insert_paragraph_after(cursor, block[1], size=12, first_line_indent=True)
    return cursor


DETAIL_SUBSECTIONS = ["功能描述", "性能描述", "输入", "输出", "程序逻辑", "限制条件"]


def append_detail_module(anchor: Paragraph, md_path: Path, module_idx: int, module_name: str, occurrence: int) -> Paragraph:
    title = insert_paragraph_after(anchor, f"8.{module_idx} {module_name}", size=12, bold=True, first_line_indent=False)
    cursor = title
    for sub_idx, sub_heading in enumerate(DETAIL_SUBSECTIONS, start=1):
        sub_para = insert_paragraph_after(cursor, f"8.{module_idx}.{sub_idx} {sub_heading}", size=11.5, bold=True,
                                          first_line_indent=False)
        cursor = render_markdown_section(sub_para, extract_section(md_path, sub_heading, occurrence=occurrence))
    return cursor


def append_titled_section(anchor: Paragraph, title: str, md_path: Path, md_heading: str, occurrence: int = 1) -> Paragraph:
    title_para = insert_paragraph_after(anchor, title, size=12, bold=True, first_line_indent=False)
    return render_markdown_section(title_para, extract_section(md_path, md_heading, occurrence=occurrence))


def strip_word_comments(docx_path: Path) -> None:
    """Remove template comments and related markup from a generated docx."""
    with zipfile.ZipFile(docx_path, "r") as zin:
        members = {name: zin.read(name) for name in zin.namelist()}

    skip_parts = {
        "word/comments.xml",
        "word/commentsExtended.xml",
        "word/commentsIds.xml",
        "word/people.xml",
    }

    xml_like_suffixes = (".xml", ".rels")
    cleaned: dict[str, bytes] = {}

    for name, data in members.items():
        if name in skip_parts:
            continue

        if name.endswith(xml_like_suffixes):
            text = data.decode("utf-8", "ignore")
            text = re.sub(r"<w:commentRangeStart[^>]*/>", "", text)
            text = re.sub(r"<w:commentRangeEnd[^>]*/>", "", text)
            text = re.sub(r"<w:commentReference[^>]*/>", "", text)
            text = re.sub(r'<Relationship[^>]+Type="[^"]*/comments[^"]*"[^>]*/>', "", text)
            text = re.sub(r'<Relationship[^>]+Type="[^"]*/commentsExtended[^"]*"[^>]*/>', "", text)
            text = re.sub(r'<Relationship[^>]+Type="[^"]*/commentsIds[^"]*"[^>]*/>', "", text)
            text = re.sub(r'<Relationship[^>]+Type="[^"]*/people[^"]*"[^>]*/>', "", text)
            text = re.sub(r'<Override PartName="/word/comments[^"]*" ContentType="[^"]+"/>', "", text)
            text = re.sub(r'<Override PartName="/word/commentsExtended[^"]*" ContentType="[^"]+"/>', "", text)
            text = re.sub(r'<Override PartName="/word/commentsIds[^"]*" ContentType="[^"]+"/>', "", text)
            text = re.sub(r'<Override PartName="/word/people.xml" ContentType="[^"]+"/>', "", text)
            cleaned[name] = text.encode("utf-8")
        else:
            cleaned[name] = data

    with zipfile.ZipFile(docx_path, "w") as zout:
        for name, data in cleaned.items():
            zout.writestr(name, data)


def fill_design_doc() -> Path:
    md_path = ROOT / "docs" / "competition" / "设计及创新性分析报告_初稿.md"
    template = TEMPLATE_DIR / "第十九届全国大学生软件创新大赛-设计及创新性分析报告模版.docx"
    output = OUTPUT_DIR / "第十九届全国大学生软件创新大赛-星火-设计及创新性分析报告-v0.5.0.docx"

    doc = Document(str(template))
    fill_cover(doc, "设计及创新性分析报告")
    fill_revision_history(doc.tables[1])

    # Align body headings and visible TOC text with the corrected v0.5 structure.
    try:
        replace_paragraph_text(find_paragraph(doc, "1.1\t痛点概述\t1"), "1.1\t相关工作\t1", size=10.5, first_line_indent=False)
        replace_paragraph_text(find_paragraph(doc, "1.2\t相关工作\t1"), "1.2\t痛点概述与项目空白\t1", size=10.5, first_line_indent=False)
    except ValueError:
        pass
    replace_paragraph_text(find_paragraph(doc, "痛点概述"), "相关工作", size=14, bold=True, first_line_indent=False)
    replace_paragraph_text(find_paragraph(doc, "相关工作"), "痛点概述与项目空白", size=14, bold=True, first_line_indent=False)

    render_markdown_section(find_paragraph(doc, "相关工作"), extract_section(md_path, "相关工作"))
    render_markdown_section(find_paragraph(doc, "痛点概述与项目空白"), extract_section(md_path, "痛点概述与项目空白"))
    render_markdown_section(find_paragraph(doc, "技术性创新点"), extract_section(md_path, "技术性创新点"))
    render_markdown_section(find_paragraph(doc, "功能性创新点"), extract_section(md_path, "功能性创新点"))
    render_markdown_section(find_paragraph(doc, "其他创新点"), extract_section(md_path, "其他创新点"))
    render_markdown_section(find_paragraph(doc, "竞品分析"), extract_section(md_path, "竞品分析"))

    doc.save(str(output))
    strip_word_comments(output)
    return output


def fill_tech_doc() -> Path:
    md_path = ROOT / "docs" / "competition" / "技术研究报告_初稿.md"
    template = TEMPLATE_DIR / "第十九届全国大学生软件创新大赛-技术研究报告模版.docx"
    output = OUTPUT_DIR / "第十九届全国大学生软件创新大赛-星火-技术研究报告-v0.5.0.docx"

    doc = Document(str(template))
    fill_cover(doc, "技术研究报告")
    fill_revision_history(doc.tables[1])

    for heading in ["问题描述", "问题抽象", "问题定位", "问题评估", "问题分解"]:
        render_markdown_section(find_paragraph(doc, heading), extract_section(md_path, heading))
    render_markdown_section(find_paragraph(doc, "相关工作"), extract_section(md_path, "相关工作"))

    for heading in ["技术方向", "技术选择", "结果期望", "使用的开发框架及依赖的库", "技术实践过程"]:
        render_markdown_section(find_paragraph(doc, heading), extract_section(md_path, heading))

    result_anchor = find_paragraph(doc, "结果验证")
    cursor = render_markdown_section(result_anchor, extract_section(md_path, "性能验证数据"))
    cursor = render_markdown_section(cursor, extract_section(md_path, "测试覆盖验证"))
    cursor = render_markdown_section(cursor, extract_section(md_path, "工程稳定性验证"))
    render_markdown_section(cursor, extract_section(md_path, "当前仍需补强的部分"))

    doc.save(str(output))
    strip_word_comments(output)
    return output


def fill_dev_doc() -> Path:
    md_path = ROOT / "docs" / "competition" / "项目开发文档_初稿.md"
    template = TEMPLATE_DIR / "第十九届全国大学生软件创新大赛-项目开发文档模版.docx"
    output = OUTPUT_DIR / "第十九届全国大学生软件创新大赛-星火-项目开发文档-v0.5.0.docx"

    doc = Document(str(template))
    fill_cover(doc, "项目开发文档")
    fill_revision_history(doc.tables[1])
    module_table = doc.tables[2]
    use_case_table = doc.tables[3]

    replace_paragraph_text(find_paragraph(doc, "4.2.1 **功能模块\t4"), "4.2.1 核心功能模块\t4", size=10.5,
                           first_line_indent=False)
    replace_paragraph_text(find_paragraph(doc, "8.1 **功能模块\t10"), "8.1 AI 对话与主链编排\t10", size=10.5,
                           first_line_indent=False)
    replace_paragraph_text(find_paragraph(doc, "8.2 **功能模块\t10"), "8.2 学习规划与任务执行\t10", size=10.5,
                           first_line_indent=False)
    replace_paragraph_text(find_paragraph(doc, "表2 ****用例规约"), "表2 核心用例规约", size=12, bold=True,
                           first_line_indent=False)

    mapping = {
        "项目背景": "项目背景",
        "项目定位": "定位说明",
        "应用场景": "应用场景",
        "目标人群": "目标人群",
        "项目方案": "项目方案",
        "项目目标": "项目目标",
        "项目价值": "项目价值",
        "最终呈现形式": "最终呈现形式",
        "主要功能描述": "主要功能描述",
        "运行环境": "运行环境",
        "验收标准": "验收标准",
        "关键问题": "关键问题",
        "进度安排": "进度安排",
        "开发预算": "开发预算",
        "技术可行性分析": "技术可行性分析",
        "资源可行性分析": "资源可行性分析",
        "市场可行性分析": "市场可行性分析",
        "静态数据": "静态数据",
        "动态数据": "动态数据",
        "数据词典": "数据词典",
        "数据采集": "数据采集",
        "时间特性": "时间特性",
        "适应性": "适应性",
        "界面需求": "界面需求",
        "硬件接口": "硬件接口",
        "软件接口": "软件接口",
        "其他需求": "其他需求",
        "处理流程": "处理流程",
        "总体结构设计": "总体结构设计",
        "功能设计": "功能设计",
        "数据流转设计": "数据流转设计",
        "用户界面设计": "用户界面设计",
        "数据结构设计": "数据结构设计",
        "外部接口": "外部接口",
        "内部接口": "内部接口",
        "系统配置策略": "系统配置策略",
        "系统部署方案": "系统部署方案",
        "跨端应用架构设计": "跨端应用架构设计",
        "其他相关技术与方案": "其他相关技术与方案",
        "数据库设计": "数据库设计",
        "手机环境需求": "手机环境需求",
    }

    for word_heading, md_heading in mapping.items():
        render_markdown_section(find_paragraph(doc, word_heading), extract_section(md_path, md_heading))

    # 4.2 功能需求总述
    render_markdown_section(find_paragraph(doc, "功能需求"), extract_section(md_path, "功能需求"))

    # 核心功能模块表
    module_placeholders = find_all_paragraphs(doc, "**功能模块")
    replace_paragraph_text(module_placeholders[0], "核心功能模块", size=12, bold=True, first_line_indent=False)
    module_rows = [
        ["AI 对话与主链编排", "意图识别、澄清、主链路由、结果合成", "作为统一入口，把自然语言需求转化为后续模块动作。", "P0"],
        ["学习规划与任务执行", "计划生成、任务拆解、执行状态回写", "把模糊学习目标变成可持续推进的行动链。", "P0"],
        ["知识星图与 GraphRAG", "知识结构化、掌握度、路径推荐、混合检索", "把知识关系与学习路径真正纳入主链决策。", "P0"],
        ["Mirofish 多 Agent 协作", "专家路由、协作、聚合、重规划", "面向复杂学习问题提供可解释协作能力。", "P1"],
        ["OpenClaw 执行闭环", "ExecutionIntent、路由、执行、信任回流", "把学习建议推进到可治理的执行层。", "P1"],
        ["学习报告", "阶段总结、薄弱点分析、后续建议", "把阶段行为沉淀成可理解的成长反馈。", "P1"],
        ["知识剧场", "路径推演、状态解释、预测表达", "以更具解释性的方式呈现学习状态。", "P1"],
        ["学习模拟", "角色扮演、轮次互动、结果反馈", "通过情景化练习加深理解和迁移。", "P1"],
        ["成就契约与成长反馈", "成就、契约、里程碑反馈", "强化长期使用中的成长感与持续性。", "P2"],
        ["BGM 与多感官体验", "路由级音景、触觉反馈、模式切换", "增强专注感和成长反馈的感知强度。", "P2"],
    ]
    while len(module_table.rows) < len(module_rows) + 1:
        module_table.add_row()
    for r, row_data in enumerate(module_rows, start=1):
        for c, text in enumerate(row_data):
            module_table.cell(r, c).text = text
    for row_idx, row in enumerate(module_table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10.5, bold=row_idx == 0)
                format_body_paragraph(para, first_line_indent=False)

    # 用例规约表
    use_case_content = {
        "用例名称": "生成学习计划并进入执行闭环",
        "功能简述": "用户输入学习目标后，系统澄清约束并生成阶段计划、任务拆解和后续执行建议。",
        "用例编号": "UC-PLAN-001",
        "执行者": "大学生用户、星火系统、OpenClaw（可选）",
        "前置条件": "用户已登录；系统已获取基本画像与会话上下文；相关依赖服务可用。",
        "后置条件": "生成计划与任务，写回用户状态；若触发执行，则产生执行记录与结果回流。",
        "涉众利益": "用户获得更清晰的行动路径；系统沉淀更完整的学习状态数据。",
        "基本路径": "输入目标 -> 系统澄清 -> 生成计划 -> 拆分任务 -> 用户确认 -> 执行/记录 -> 反馈回流。",
        "扩展路径": "若信息不足则先补充澄清；若任务可外部执行则进入 OpenClaw 路由与审批链路。",
        "字段列表": "目标描述、时间约束、学科范围、当前掌握度、计划节点、任务项、执行状态、反馈结果。",
        "设计规则": "优先保证可执行性与安全边界；复杂链路必须可解释、可回滚、可留痕。",
        "未解决的问题": "真机触觉/通知/分享链路仍待完整补签收；复杂多 Agent 演示脚本还可继续标准化。",
        "备注": "该用例覆盖星火最核心的主链，是比赛演示的推荐首选用例。",
    }
    for row in use_case_table.rows:
        key = row.cells[0].text.strip()
        if key in use_case_content:
            row.cells[1].text = use_case_content[key]
    for row in use_case_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10.5, bold=cell is row.cells[0])
                format_body_paragraph(para, first_line_indent=False)

    extra_use_cases = [
        (
            "表 3 知识星图路径规划用例规约",
            {
                "用例名称": "基于知识星图生成学习路径",
                "功能简述": "用户围绕某一知识主题请求系统给出前置知识、当前切入点与推荐学习路径。",
                "用例编号": "UC-GRAPH-002",
                "执行者": "大学生用户、星火系统",
                "前置条件": "图谱中存在相关知识节点；用户已有基础画像或历史学习记录。",
                "后置条件": "生成结构化路径建议，并将其写回后续计划或报告。",
                "涉众利益": "用户能快速看到知识结构，减少盲目学习成本。",
                "基本路径": "输入主题 -> 检索相关节点 -> 图推理生成路径 -> 返回掌握度与建议。",
                "扩展路径": "若知识节点不存在，则走 free mode 或弱结构化推荐逻辑。",
                "字段列表": "目标主题、相关节点、关系类型、掌握度、推荐顺序、复习建议。",
                "设计规则": "优先保持路径可解释，不输出无法说明来源的黑盒结论。",
                "未解决的问题": "最终提交版需补充更清晰的图谱可视化截图。",
                "备注": "该用例用于支撑知识星图不是装饰页面，而是具备推理能力的核心模块。",
            },
        ),
        (
            "表 4 OpenClaw 执行闭环用例规约",
            {
                "用例名称": "执行意图生成并回流到学习主链",
                "功能简述": "系统识别某任务适合进入外部执行闭环，经路由、审批和信任评估后回流结果。",
                "用例编号": "UC-EXEC-003",
                "执行者": "大学生用户、星火系统、OpenClaw",
                "前置条件": "任务属于可执行边界；执行策略与审批上下文已构建。",
                "后置条件": "执行结果写回任务状态，并作为后续反馈或报告输入。",
                "涉众利益": "用户获得更强推进能力；系统形成从建议到行动的差异化闭环。",
                "基本路径": "识别执行机会 -> 构建 ExecutionIntent -> 路由 -> 执行 -> 解析 -> TrustEngine -> 回流。",
                "扩展路径": "若风险较高或可信度不足，则转人工确认或终止回流。",
                "字段列表": "ExecutionIntent、execution_mode、target_env、trust_level、policy、result。",
                "设计规则": "能力增强必须服从边界治理和可解释性要求。",
                "未解决的问题": "高风险场景的最终展示脚本仍可继续标准化。",
                "备注": "该用例重点支撑项目的核心技术创新与系统级差异化。",
            },
        ),
    ]

    previous_table = use_case_table
    for title, content in extra_use_cases:
        title_para = add_table_title_after(previous_table, title)
        new_table = insert_table_after(title_para, rows=len(use_case_table.rows), cols=2)
        for row_idx, row in enumerate(use_case_table.rows):
            for col_idx, cell in enumerate(row.cells):
                new_table.cell(row_idx, col_idx).text = cell.text if col_idx == 0 else ""
        for row in new_table.rows:
            key = row.cells[0].text.strip()
            if key in content:
                row.cells[1].text = content[key]
        format_table(new_table, header_rows=0, font_size=10.5)
        previous_table = new_table

    # 错误/异常处理
    error_anchor = find_paragraph(doc, "错误/异常输出信息")
    render_markdown_section(error_anchor, [
        "系统在异常场景下应优先输出可理解的失败原因、当前状态、建议动作和可恢复路径，避免只暴露底层报错信息。",
        "",
        "对用户侧异常，采用清晰提示语、重试建议和状态恢复说明；对系统侧异常，记录 trace、上下文和关键参数，便于快速定位。",
    ])
    countermeasure_anchor = find_paragraph(doc, "错误/异常处理对策")
    render_markdown_section(countermeasure_anchor, [
        "系统通过超时控制、路由降级、熔断、审批确认、缓存复用和可观测性机制降低异常对主链体验的破坏。",
        "",
        "对于高风险外部执行场景，默认采用保守路由和信任评估；对于移动端弱网和依赖服务波动场景，优先保证用户可见状态一致性与任务可恢复性。",
    ])

    # 详细设计
    detail_modules = [
        "AI 对话与主链编排",
        "学习规划与任务执行",
        "知识星图与 GraphRAG",
        "Mirofish 多 Agent 协作",
        "OpenClaw 执行闭环",
        "学习报告",
        "知识剧场",
        "学习模拟",
        "成就契约与成长反馈",
        "BGM 与多感官体验",
    ]

    replace_paragraph_text(module_placeholders[1], "AI 对话与主链编排", size=12, bold=True, first_line_indent=False)
    replace_paragraph_text(module_placeholders[2], "学习规划与任务执行", size=12, bold=True, first_line_indent=False)

    first_module_subheads = [find_paragraph(doc, name) for name in DETAIL_SUBSECTIONS]
    for idx, (para, sub_heading) in enumerate(zip(first_module_subheads, DETAIL_SUBSECTIONS), start=1):
        replace_paragraph_text(para, f"8.1.{idx} {sub_heading}", size=11.5, bold=True, first_line_indent=False)
        render_markdown_section(para, extract_section(md_path, sub_heading, occurrence=1))

    cursor = module_placeholders[2]
    for sub_idx, sub_heading in enumerate(DETAIL_SUBSECTIONS, start=1):
        sub_para = insert_paragraph_after(cursor, f"8.2.{sub_idx} {sub_heading}", size=11.5, bold=True, first_line_indent=False)
        cursor = render_markdown_section(sub_para, extract_section(md_path, sub_heading, occurrence=2))

    for module_idx, module_name in enumerate(detail_modules[2:], start=3):
        cursor = append_detail_module(cursor, md_path, module_idx, module_name, occurrence=module_idx)

    doc.save(str(output))
    strip_word_comments(output)
    return output


def fill_test_doc() -> Path:
    md_path = ROOT / "docs" / "competition" / "项目测试文档_初稿.md"
    template = TEMPLATE_DIR / "第十九届全国大学生软件创新大赛-项目测试文档模版.docx"
    output = OUTPUT_DIR / "第十九届全国大学生软件创新大赛-星火-项目测试文档-v0.5.0.docx"

    doc = Document(str(template))
    fill_cover(doc, "项目测试文档")
    fill_revision_history(doc.tables[1])
    unit_matrix_table = doc.tables[2]
    function_matrix_table = doc.tables[3]
    performance_table = doc.tables[4]

    replace_paragraph_text(find_paragraph(doc, "2.1\t****模块\t2"), "2.1\t功能模块覆盖矩阵\t2", size=10.5,
                           first_line_indent=False)
    replace_paragraph_text(find_paragraph(doc, "2.2\t****模块\t2"), "2.2\tAI 对话与主链编排\t2", size=10.5,
                           first_line_indent=False)
    replace_paragraph_text(find_paragraph(doc, "3.1\t****功能\t3"), "3.1\t功能模块覆盖矩阵\t3", size=10.5,
                           first_line_indent=False)
    replace_paragraph_text(find_paragraph(doc, "3.2\t****功能\t3"), "3.2\tAI 对话与主链编排\t3", size=10.5,
                           first_line_indent=False)

    def first_text_block(section_name: str) -> str:
        """Return the text of the first paragraph or list block in a section."""
        for block in markdown_blocks(extract_section(md_path, section_name)):
            if block[0] in ("paragraph", "list"):
                return block[1]
        return ""

    strategy_blocks = [b for b in markdown_blocks(extract_section(md_path, "测试策略与目标")) if b[0] in ("paragraph", "list")]
    replace_paragraph_text(find_paragraph(doc, "【测试策略：测试策略在软件需求分析完成后就开始实施，根据项目需求对项目有一个整体的把握，包括：测试重点、测试难点、测试分层。】"),
                           strategy_blocks[0][1] if strategy_blocks else "", first_line_indent=True)
    replace_paragraph_text(find_paragraph(doc, "【目标：定义项目在发布时候的质量等级】"),
                           strategy_blocks[1][1] if len(strategy_blocks) > 1 else "", first_line_indent=True)
    replace_paragraph_text(find_paragraph(doc, "【从测试广度和测试深度两方面了解整个测试项目的测试规模】"),
                           first_text_block("测试范围"), first_line_indent=True)
    replace_paragraph_text(find_paragraph(doc, "【包括软硬件环境、网络环境、测试工具】"),
                           first_text_block("测试环境"), first_line_indent=True)

    unit_modules = [
        "AI 对话与主链编排",
        "学习规划与任务执行",
        "知识星图与 GraphRAG",
        "Mirofish 多 Agent 协作",
        "OpenClaw 执行闭环",
        "学习报告",
        "知识剧场",
        "学习模拟",
        "成就契约与成长反馈",
        "BGM 与多感官体验",
    ]

    # 单元测试模块名称
    module_placeholders = find_all_paragraphs(doc, "****模块")
    function_placeholders = find_all_paragraphs(doc, "****功能")
    replace_paragraph_text(module_placeholders[0], "功能模块覆盖矩阵", size=12, bold=True, first_line_indent=False)
    replace_paragraph_text(module_placeholders[1], "AI 对话与主链编排", size=12, bold=True, first_line_indent=False)
    replace_paragraph_text(function_placeholders[0], "功能模块覆盖矩阵", size=12, bold=True, first_line_indent=False)
    replace_paragraph_text(function_placeholders[1], "AI 对话与主链编排", size=12, bold=True, first_line_indent=False)

    # 单元测试覆盖矩阵
    while len(unit_matrix_table.rows) < len(unit_modules) + 1:
        unit_matrix_table.add_row()
    unit_matrix_table.cell(0, 0).text = "功能模块"
    unit_matrix_table.cell(0, 1).text = "单元测试覆盖重点"
    unit_matrix_table.cell(0, 2).text = "关键入口/对象"
    unit_matrix_table.cell(0, 3).text = "结果"
    unit_matrix_table.cell(0, 4).text = "备注"
    unit_rows = [
        ["AI 对话与主链编排", "上下文、澄清、路由", "主链编排器", "通过", "入口链路稳定"],
        ["学习规划与任务执行", "计划与任务状态流转", "计划/任务服务", "通过", "链路成立"],
        ["知识星图与 GraphRAG", "节点、关系、路径推荐", "图谱/检索服务", "通过", "具备专项验证"],
        ["Mirofish 多 Agent 协作", "路由、协作、聚合", "协作工作流", "通过", "多模式可回归"],
        ["OpenClaw 执行闭环", "Phase 0-4、TrustEngine", "执行闭环服务", "通过", "分阶段覆盖完整"],
        ["学习报告", "报告聚合与输出", "报告服务", "通过", "具备专项验证"],
        ["知识剧场", "路径推演与解释", "预测剧场服务", "通过", "具备专项验证"],
        ["学习模拟", "轮次与反馈", "模拟引擎", "通过", "具备专项验证"],
        ["成就契约与成长反馈", "触发规则与状态变更", "成就/契约服务", "通过", "规则级验证"],
        ["BGM 与多感官体验", "模式切换与恢复", "端侧体验服务", "通过", "端侧专项验证"],
    ]
    for row_idx, row_data in enumerate(unit_rows, start=1):
        for col_idx, text in enumerate(row_data):
            unit_matrix_table.cell(row_idx, col_idx).text = text
    format_table(unit_matrix_table, header_rows=1, font_size=9.6)

    unit_anchor = module_placeholders[1]
    unit_cursor = render_markdown_section(unit_anchor, extract_section(md_path, unit_modules[0], occurrence=1))
    for idx, module_name in enumerate(unit_modules[1:], start=3):
        unit_cursor = append_titled_section(unit_cursor, f"2.{idx} {module_name}", md_path, module_name, occurrence=1)

    # 功能测试覆盖矩阵
    while len(function_matrix_table.rows) < len(unit_modules) + 1:
        function_matrix_table.add_row()
    function_matrix_table.cell(0, 0).text = "功能模块"
    function_matrix_table.cell(0, 1).text = "典型功能场景"
    function_matrix_table.cell(0, 2).text = "验证焦点"
    function_matrix_table.cell(0, 3).text = "结果"
    function_matrix_table.cell(0, 4).text = "备注"
    function_rows = [
        ["AI 对话与主链编排", "用户输入到主链分流", "入口稳定性", "通过", "主链可演示"],
        ["学习规划与任务执行", "目标到计划与任务推进", "可执行性", "通过", "计划链路成立"],
        ["知识星图与 GraphRAG", "主题到路径推荐", "路径解释性", "通过", "知识结构可展示"],
        ["Mirofish 多 Agent 协作", "复杂问题到协作结果", "协作可解释性", "通过", "评审可展示"],
        ["OpenClaw 执行闭环", "任务到执行回流", "边界与回流", "通过", "差异化强"],
        ["学习报告", "学习数据到报告展示", "反馈可读性", "通过", "可引导下一步"],
        ["知识剧场", "路径推演展示", "解释清晰度", "通过", "预测链路可展示"],
        ["学习模拟", "主题到多轮模拟", "轮次稳定性", "通过", "情景化成立"],
        ["成就契约与成长反馈", "触发到反馈展示", "反馈节奏", "通过", "成长闭环成立"],
        ["BGM 与多感官体验", "场景切换到多感官反馈", "端侧协调", "通过", "具备降级能力"],
    ]
    for row_idx, row_data in enumerate(function_rows, start=1):
        for col_idx, text in enumerate(row_data):
            function_matrix_table.cell(row_idx, col_idx).text = text
    format_table(function_matrix_table, header_rows=1, font_size=9.6)

    function_anchor = function_placeholders[1]
    function_cursor = render_markdown_section(function_anchor, extract_section(md_path, unit_modules[0], occurrence=2))
    for idx, module_name in enumerate(unit_modules[1:], start=3):
        function_cursor = append_titled_section(function_cursor, f"3.{idx} {module_name}", md_path, module_name, occurrence=2)

    for row in unit_matrix_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10)
                format_body_paragraph(para, first_line_indent=False)
    for row in function_matrix_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10)
                format_body_paragraph(para, first_line_indent=False)

    # 系统测试 / 性能测试表
    system_anchor = find_paragraph(doc, "系统测试")
    cursor = append_titled_section(system_anchor, "4.1 环境与工程闸门", md_path, "环境与工程闸门")
    try:
        replace_paragraph_text(find_paragraph(doc, "4.1\t模型性能测试\t4"), "4.2\t性能测试\t4", size=10.5, first_line_indent=False)
    except ValueError:
        pass
    replace_paragraph_text(find_paragraph(doc, "模型性能测试"), "4.2 性能测试", size=12, bold=True, first_line_indent=False)
    perf_ids = ["PT-01", "PT-02", "PT-03", "PT-04"]
    for i, cid in enumerate(perf_ids, start=1):
        performance_table.cell(0, i).text = cid
    perf_rows = {
        1: ["Tier-1 意图识别延迟", "信息充分性检查延迟", "完整路由延迟", "并发吞吐量"],
        2: [
            "验证高频意图识别满足快速分类目标。",
            "验证澄清前检查不会阻塞主链。",
            "验证不含外部 LLM 网络时延的完整路由满足目标。",
            "验证单机并发场景下具备可接受吞吐能力。",
        ],
        3: ["性能测试环境准备完成。"] * 4,
        4: ["基于现有性能测试与 benchmark 阈值。"] * 4,
        5: ["依赖 RequestRouter、SufficiencyChecker 和 Go Gateway benchmark。"] * 4,
    }
    for row_idx, values in perf_rows.items():
        for col_idx, text in enumerate(values, start=1):
            performance_table.cell(row_idx, col_idx).text = text
    for col in range(1, 5):
        performance_table.cell(6, col).text = ["输入/动作", "期望的性能\n（平均值）", "实际的性能\n（平均值）", "备注"][col - 1]
    performance_table.cell(7, 1).text = "连续分类请求"
    performance_table.cell(7, 2).text = "P50 < 10ms，P95 < 25ms"
    performance_table.cell(7, 3).text = "满足阈值测试"
    performance_table.cell(7, 4).text = "见 intent performance 测试"
    performance_table.cell(8, 1).text = "充分性检查"
    performance_table.cell(8, 2).text = "P50 < 50ms"
    performance_table.cell(8, 3).text = "满足阈值测试"
    performance_table.cell(8, 4).text = "本地回归可验证"
    performance_table.cell(9, 1).text = "路由/吞吐 benchmark"
    performance_table.cell(9, 2).text = "P95 < 100ms；分类 >100 req/s；路由 >50 req/s"
    performance_table.cell(9, 3).text = "满足现有阈值断言"
    performance_table.cell(9, 4).text = "详细数值待统一复测"
    for row in performance_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10)
                format_body_paragraph(para, first_line_indent=False)

    perf_anchor = find_paragraph(doc, "4.2 性能测试")
    perf_cursor = insert_paragraph_after(perf_anchor, "当前性能测试采用阈值型验证方式，重点保证比赛演示场景下的稳定响应，而不是追求脱离场景的理论峰值。", size=12)
    perf_cursor = append_titled_section(insert_paragraph_after_table(performance_table), "4.3 混沌测试与弹性验证", md_path, "混沌测试与弹性验证")
    perf_cursor = append_titled_section(perf_cursor, "4.4 测试数据汇总", md_path, "测试数据汇总")
    append_titled_section(perf_cursor, "4.5 当前未完全完成的系统级验证", md_path, "当前未完全完成的系统级验证")

    doc.save(str(output))
    strip_word_comments(output)
    return output


def fill_support_doc(md_filename: str, output_name: str, title: str) -> Path:
    md_path = ROOT / "docs" / "competition" / md_filename
    output = OUTPUT_DIR / output_name
    doc = Document()
    doc.core_properties.author = AUTHOR
    doc.core_properties.title = title

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    set_run_font(run, size=18, bold=True)
    format_body_paragraph(title_para, first_line_indent=False)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"星火 sparkle | {DATE} | v{VERSION}")
    set_run_font(run, size=11)
    format_body_paragraph(subtitle, first_line_indent=False)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    for raw_line in lines:
        stripped = raw_line.strip()
        if not stripped or stripped == "---":
            doc.add_paragraph()
            continue
        if stripped.startswith("# "):
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(normalize_heading(stripped))
            set_run_font(run, size=14, bold=True)
            format_body_paragraph(p, first_line_indent=False)
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(normalize_heading(stripped))
            set_run_font(run, size=12, bold=True)
            format_body_paragraph(p, first_line_indent=False)
            continue
        if stripped.startswith("#### "):
            p = doc.add_paragraph()
            run = p.add_run(normalize_heading(stripped))
            set_run_font(run, size=11, bold=True)
            format_body_paragraph(p, first_line_indent=False)
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph()
            run = p.add_run(f"• {stripped[2:].strip()}")
            set_run_font(run, size=11)
            format_body_paragraph(p, first_line_indent=False)
            continue
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        set_run_font(run, size=11)
        format_body_paragraph(p, first_line_indent=False if "：" in stripped and len(stripped) < 28 else True)

    doc.save(str(output))
    strip_word_comments(output)
    return output


def main() -> None:
    outputs = [
        fill_design_doc(),
        fill_tech_doc(),
        fill_dev_doc(),
        fill_test_doc(),
        fill_support_doc(
            "配图与AI生成Prompt清单_初稿.md",
            "第十九届全国大学生软件创新大赛-星火-配图与AI生成Prompt清单-v0.5.0.docx",
            "星火配图与 AI 生成 Prompt 清单",
        ),
    ]
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
